#!/usr/bin/env python3
"""
Forge - ZIM File Creator for Hermit

Create custom ZIM knowledge bases from your documents.
Supports: PDF, DOCX, TXT, Markdown, HTML, EPUB

Usage:
    forge              # Launch GUI
    forge --cli        # Interactive CLI mode
    forge /path/to/docs --output my_knowledge.zim  # CLI batch mode
"""

import os
import sys
import argparse
import hashlib
import html
from pathlib import Path
from typing import List, Dict, Optional, Generator
from datetime import datetime
import re

# Document Parsing Libraries (optional, checked at runtime)
PARSERS_AVAILABLE = {
    'pdf': False,
    'docx': False,
    'epub': False,
    'markdown': False
}

try:
    import pypdf
    PARSERS_AVAILABLE['pdf'] = True
except ImportError:
    pass

try:
    import docx
    PARSERS_AVAILABLE['docx'] = True
except ImportError:
    pass

try:
    import ebooklib
    from ebooklib import epub
    PARSERS_AVAILABLE['epub'] = True
except ImportError:
    pass

try:
    import markdown
    PARSERS_AVAILABLE['markdown'] = True
except ImportError:
    pass

# ZIM Writer
try:
    from libzim.writer import Creator, Item, StringProvider, FileProvider, Hint
    LIBZIM_AVAILABLE = True
except ImportError:
    LIBZIM_AVAILABLE = False


class Document:
    """Represents a parsed document."""
    
    def __init__(self, title: str, content: str, source_path: str, 
                 doc_type: str = "article", metadata: Dict = None):
        self.title = title
        self.content = content  # Plain text
        self.source_path = source_path
        self.doc_type = doc_type
        self.metadata = metadata or {}
        self.word_count = len(content.split())
        
        # Generate unique path for ZIM
        self.zim_path = self._generate_path()
    
    def _generate_path(self) -> str:
        """Generate a URL-safe path for this document."""
        # Clean title for URL
        safe_title = re.sub(r'[^\w\s-]', '', self.title)
        safe_title = re.sub(r'\s+', '_', safe_title).strip('_')
        
        # Add hash for uniqueness
        path_hash = hashlib.md5(self.source_path.encode()).hexdigest()[:8]
        
        return f"A/{safe_title}_{path_hash}"
    
    def to_html(self) -> str:
        """Convert document to HTML for ZIM storage."""
        # Escape content and convert newlines to paragraphs
        escaped = html.escape(self.content)
        paragraphs = escaped.split('\n\n')
        
        html_paras = '\n'.join(f'<p>{p}</p>' for p in paragraphs if p.strip())
        
        return f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{html.escape(self.title)}</title>
</head>
<body>
    <h1>{html.escape(self.title)}</h1>
    <p class="metadata">Source: {html.escape(os.path.basename(self.source_path))}</p>
    <hr>
    {html_paras}
</body>
</html>
"""


class DocumentParser:
    """Parse various document formats into plain text."""
    
    @staticmethod
    def parse_file(file_path: str) -> Optional[Document]:
        """Parse a file and return a Document object."""
        path = Path(file_path)
        
        if not path.exists():
            print(f"[WARN] File not found: {file_path}")
            return None
        
        ext = path.suffix.lower()
        title = path.stem
        
        try:
            if ext == '.txt':
                content = DocumentParser._parse_txt(path)
            elif ext == '.md':
                content = DocumentParser._parse_markdown(path)
            elif ext == '.pdf':
                content = DocumentParser._parse_pdf(path)
            elif ext == '.docx':
                content = DocumentParser._parse_docx(path)
            elif ext in ['.html', '.htm']:
                content = DocumentParser._parse_html(path)
            elif ext == '.epub':
                content = DocumentParser._parse_epub(path)
            else:
                print(f"[WARN] Unsupported format: {ext}")
                return None
            
            if not content or len(content.strip()) < 10:
                print(f"[WARN] Empty or too short: {file_path}")
                return None
            
            return Document(
                title=title,
                content=content,
                source_path=str(path.absolute()),
                doc_type="article",
                metadata={"format": ext, "size": path.stat().st_size}
            )
            
        except Exception as e:
            print(f"[ERROR] Failed to parse {file_path}: {e}")
            return None
    
    @staticmethod
    def _parse_txt(path: Path) -> str:
        """Parse plain text file."""
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    
    @staticmethod
    def _parse_markdown(path: Path) -> str:
        """Parse Markdown file to plain text."""
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        
        if PARSERS_AVAILABLE['markdown']:
            # Convert MD to HTML then strip tags
            md = markdown.markdown(content)
            return re.sub(r'<[^>]+>', '', md)
        else:
            # Basic: just remove markdown syntax
            content = re.sub(r'^#{1,6}\s+', '', content, flags=re.MULTILINE)
            content = re.sub(r'\*\*([^*]+)\*\*', r'\1', content)
            content = re.sub(r'\*([^*]+)\*', r'\1', content)
            content = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', content)
            return content
    
    @staticmethod
    def _parse_pdf(path: Path) -> str:
        """Parse PDF file."""
        if not PARSERS_AVAILABLE['pdf']:
            raise ImportError("pypdf not installed. Run: pip install pypdf")
        
        reader = pypdf.PdfReader(str(path))
        text_parts = []
        
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        
        return '\n\n'.join(text_parts)
    
    @staticmethod
    def _parse_docx(path: Path) -> str:
        """Parse Word DOCX file."""
        if not PARSERS_AVAILABLE['docx']:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
        
        doc = docx.Document(str(path))
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        return '\n\n'.join(paragraphs)
    
    @staticmethod
    def _parse_html(path: Path) -> str:
        """Parse HTML file to plain text."""
        from bs4 import BeautifulSoup
        
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            soup = BeautifulSoup(f.read(), 'html.parser')
        
        # Remove script and style elements
        for tag in soup(['script', 'style', 'nav', 'header', 'footer']):
            tag.decompose()
        
        return soup.get_text(separator='\n\n')
    
    @staticmethod
    def _parse_epub(path: Path) -> str:
        """Parse EPUB ebook."""
        if not PARSERS_AVAILABLE['epub']:
            raise ImportError("ebooklib not installed. Run: pip install ebooklib")
        
        from bs4 import BeautifulSoup
        
        book = epub.read_epub(str(path))
        text_parts = []
        
        for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
            soup = BeautifulSoup(item.get_content(), 'html.parser')
            text = soup.get_text(separator='\n')
            if text.strip():
                text_parts.append(text)
        
        return '\n\n'.join(text_parts)


class ZIMItem(Item):
    """A single item (article) in the ZIM file."""
    
    def __init__(self, doc: Document):
        super().__init__()
        self.doc = doc
        self._html = doc.to_html().encode('utf-8')
    
    def get_path(self) -> str:
        return self.doc.zim_path
    
    def get_title(self) -> str:
        return self.doc.title
    
    def get_mimetype(self) -> str:
        return "text/html"
    
    def get_contentprovider(self):
        return StringProvider(self._html.decode('utf-8'))
    
    def get_hints(self):
        return {Hint.FRONT_ARTICLE: True}


class ZIMHomePage(Item):
    """The main index page for the ZIM file."""
    
    def __init__(self, title: str, docs: List[Document]):
        super().__init__()
        self.title = title
        self.docs = docs
    
    def get_path(self) -> str:
        return "A/index"
    
    def get_title(self) -> str:
        return self.title
    
    def get_mimetype(self) -> str:
        return "text/html"
    
    def get_contentprovider(self):
        # Build index HTML
        doc_list = '\n'.join(
            f'<li><a href="{doc.zim_path}">{html.escape(doc.title)}</a> '
            f'<small>({doc.word_count} words)</small></li>'
            for doc in self.docs
        )
        
        html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{html.escape(self.title)}</title>
    <style>
        body {{ font-family: sans-serif; max-width: 800px; margin: 2rem auto; padding: 1rem; }}
        h1 {{ color: #333; }}
        ul {{ line-height: 1.8; }}
        a {{ color: #0066cc; text-decoration: none; }}
        a:hover {{ text-decoration: underline; }}
        small {{ color: #666; }}
    </style>
</head>
<body>
    <h1>📚 {html.escape(self.title)}</h1>
    <p>This knowledge base contains <strong>{len(self.docs)}</strong> documents.</p>
    <p>Created with <a href="https://github.com/imDelivered/Hermit-AI">Hermit Forge</a></p>
    <hr>
    <h2>Documents</h2>
    <ul>
        {doc_list}
    </ul>
</body>
</html>
"""
        return StringProvider(html_content)
    
    def get_hints(self):
        return {Hint.FRONT_ARTICLE: True}


class ZIMCreator:
    """Create ZIM files from documents."""
    
    def __init__(self, output_path: str, title: str = "My Knowledge Base"):
        if not LIBZIM_AVAILABLE:
            raise ImportError(
                "libzim not available. Install with: pip install libzim\n"
                "Or system: sudo apt install python3-libzim"
            )
        
        self.output_path = output_path
        self.title = title
        self.documents: List[Document] = []
    
    def add_document(self, doc: Document):
        """Add a parsed document."""
        self.documents.append(doc)
    
    def add_directory(self, dir_path: str, recursive: bool = True) -> int:
        """Scan a directory and add all supported files."""
        path = Path(dir_path)
        
        if not path.is_dir():
            raise ValueError(f"Not a directory: {dir_path}")
        
        extensions = {'.txt', '.md', '.pdf', '.docx', '.html', '.htm', '.epub'}
        count = 0
        
        if recursive:
            files = path.rglob('*')
        else:
            files = path.glob('*')
        
        for file_path in files:
            if file_path.suffix.lower() in extensions:
                doc = DocumentParser.parse_file(str(file_path))
                if doc:
                    self.add_document(doc)
                    count += 1
                    print(f"[OK] Added: {file_path.name}")
        
        return count
    
    def create(self, compression: str = "zstd") -> str:
        """Create the ZIM file."""
        if not self.documents:
            raise ValueError("No documents to add. Add documents first.")
        
        print(f"\n[INFO] Creating ZIM file: {self.output_path}")
        print(f"[INFO] Documents: {len(self.documents)}")
        
        # Create and configure BEFORE entering context manager
        creator = Creator(self.output_path)
        creator.config_indexing(True, "en")
        creator.set_mainpath("A/index")
        
        # Enter context manager to start writing
        with creator:
            # Add homepage
            homepage = ZIMHomePage(self.title, self.documents)
            creator.add_item(homepage)
            
            # Add all documents
            for i, doc in enumerate(self.documents):
                item = ZIMItem(doc)
                creator.add_item(item)
                
                if (i + 1) % 10 == 0:
                    print(f"[PROGRESS] Added {i + 1}/{len(self.documents)} documents")
        
        # Get file size
        size_mb = os.path.getsize(self.output_path) / (1024 * 1024)
        print(f"\n[SUCCESS] Created: {self.output_path} ({size_mb:.1f} MB)")
        
        return self.output_path


# =============================================================================
# GUI Interface
# =============================================================================

class ForgeGUI:
    """Tkinter GUI for Forge ZIM creator."""
    
    def __init__(self):
        try:
            import tkinter as tk
            from tkinter import ttk, filedialog, messagebox, scrolledtext
        except ImportError:
            raise RuntimeError("tkinter not available. Install: sudo apt install python3-tk")
        
        self.tk = tk
        self.ttk = ttk
        self.filedialog = filedialog
        self.messagebox = messagebox
        self.scrolledtext = scrolledtext
        
        self.root = tk.Tk()
        self.root.title("Hermit Forge - ZIM Creator")
        self.root.geometry("700x600")
        
        # File list
        self.files: List[str] = []
        
        self._setup_ui()
    
    def _setup_ui(self):
        """Setup the GUI components."""
        tk = self.tk
        ttk = self.ttk
        
        # Title
        title_frame = ttk.Frame(self.root)
        title_frame.pack(fill=tk.X, padx=20, pady=10)
        
        ttk.Label(title_frame, text="🔨 Hermit Forge", font=("Helvetica", 20, "bold")).pack(side=tk.LEFT)
        ttk.Label(title_frame, text="Create ZIM knowledge bases", font=("Helvetica", 10)).pack(side=tk.LEFT, padx=10)
        
        # Settings Frame
        settings_frame = ttk.LabelFrame(self.root, text="Settings", padding=10)
        settings_frame.pack(fill=tk.X, padx=20, pady=5)
        
        # ZIM Title
        ttk.Label(settings_frame, text="Knowledge Base Title:").grid(row=0, column=0, sticky=tk.W)
        self.title_var = tk.StringVar(value="My Knowledge Base")
        ttk.Entry(settings_frame, textvariable=self.title_var, width=40).grid(row=0, column=1, padx=5)
        
        # Output path
        ttk.Label(settings_frame, text="Output File:").grid(row=1, column=0, sticky=tk.W, pady=5)
        self.output_var = tk.StringVar(value="knowledge.zim")
        ttk.Entry(settings_frame, textvariable=self.output_var, width=40).grid(row=1, column=1, padx=5)
        ttk.Button(settings_frame, text="Browse...", command=self._browse_output).grid(row=1, column=2)
        
        # Files Frame
        files_frame = ttk.LabelFrame(self.root, text="Source Documents", padding=10)
        files_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=5)
        
        # File list
        list_frame = ttk.Frame(files_frame)
        list_frame.pack(fill=tk.BOTH, expand=True)
        
        scrollbar = ttk.Scrollbar(list_frame)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        self.file_listbox = tk.Listbox(list_frame, yscrollcommand=scrollbar.set, height=12)
        self.file_listbox.pack(fill=tk.BOTH, expand=True)
        scrollbar.config(command=self.file_listbox.yview)
        
        # File buttons
        btn_frame = ttk.Frame(files_frame)
        btn_frame.pack(fill=tk.X, pady=5)
        
        ttk.Button(btn_frame, text="📄 Add Files", command=self._add_files).pack(side=tk.LEFT, padx=2)
        ttk.Button(btn_frame, text="📁 Add Folder", command=self._add_folder).pack(side=tk.LEFT, padx=2)
        ttk.Button(btn_frame, text="❌ Remove Selected", command=self._remove_selected).pack(side=tk.LEFT, padx=2)
        ttk.Button(btn_frame, text="🗑️ Clear All", command=self._clear_all).pack(side=tk.LEFT, padx=2)
        
        # Supported formats label
        formats_text = "Supported: TXT, MD, PDF*, DOCX*, HTML, EPUB*"
        if not PARSERS_AVAILABLE['pdf']:
            formats_text += "\n* PDF requires: pip install pypdf"
        if not PARSERS_AVAILABLE['docx']:
            formats_text += "\n* DOCX requires: pip install python-docx"
        if not PARSERS_AVAILABLE['epub']:
            formats_text += "\n* EPUB requires: pip install ebooklib"
        
        ttk.Label(files_frame, text=formats_text, font=("Helvetica", 9), foreground="gray").pack(anchor=tk.W)
        
        # Status/Log
        log_frame = ttk.LabelFrame(self.root, text="Log", padding=5)
        log_frame.pack(fill=tk.X, padx=20, pady=5)
        
        self.log_text = self.scrolledtext.ScrolledText(log_frame, height=5, state=tk.DISABLED)
        self.log_text.pack(fill=tk.X)
        
        # Create button
        action_frame = ttk.Frame(self.root)
        action_frame.pack(fill=tk.X, padx=20, pady=10)
        
        self.status_var = tk.StringVar(value="Ready")
        ttk.Label(action_frame, textvariable=self.status_var, font=("Helvetica", 10)).pack(side=tk.LEFT)
        
        ttk.Button(action_frame, text="🔨 Create ZIM", command=self._create_zim).pack(side=tk.RIGHT)
    
    def _log(self, message: str):
        """Add message to log."""
        self.log_text.config(state=self.tk.NORMAL)
        self.log_text.insert(self.tk.END, message + "\n")
        self.log_text.see(self.tk.END)
        self.log_text.config(state=self.tk.DISABLED)
        self.root.update_idletasks()
    
    def _browse_output(self):
        """Browse for output file location."""
        path = self.filedialog.asksaveasfilename(
            defaultextension=".zim",
            filetypes=[("ZIM files", "*.zim"), ("All files", "*.*")]
        )
        if path:
            self.output_var.set(path)
    
    def _add_files(self):
        """Add individual files."""
        files = self.filedialog.askopenfilenames(
            filetypes=[
                ("All supported", "*.txt *.md *.pdf *.docx *.html *.htm *.epub"),
                ("Text files", "*.txt"),
                ("Markdown", "*.md"),
                ("PDF", "*.pdf"),
                ("Word", "*.docx"),
                ("HTML", "*.html *.htm"),
                ("EPUB", "*.epub"),
            ]
        )
        for f in files:
            if f not in self.files:
                self.files.append(f)
                self.file_listbox.insert(self.tk.END, os.path.basename(f))
        
        self._update_status()
    
    def _add_folder(self):
        """Add all files from a folder."""
        folder = self.filedialog.askdirectory()
        if folder:
            extensions = {'.txt', '.md', '.pdf', '.docx', '.html', '.htm', '.epub'}
            count = 0
            for root, _, files in os.walk(folder):
                for file in files:
                    if Path(file).suffix.lower() in extensions:
                        full_path = os.path.join(root, file)
                        if full_path not in self.files:
                            self.files.append(full_path)
                            self.file_listbox.insert(self.tk.END, file)
                            count += 1
            
            self._log(f"Added {count} files from {folder}")
            self._update_status()
    
    def _remove_selected(self):
        """Remove selected files."""
        selection = self.file_listbox.curselection()
        for i in reversed(selection):
            self.file_listbox.delete(i)
            del self.files[i]
        self._update_status()
    
    def _clear_all(self):
        """Clear all files."""
        self.files.clear()
        self.file_listbox.delete(0, self.tk.END)
        self._update_status()
    
    def _update_status(self):
        """Update status label."""
        self.status_var.set(f"{len(self.files)} files selected")
    
    def _create_zim(self):
        """Create the ZIM file."""
        if not self.files:
            self.messagebox.showwarning("No Files", "Please add some documents first.")
            return
        
        if not LIBZIM_AVAILABLE:
            self.messagebox.showerror(
                "Missing Dependency",
                "libzim is not installed.\n\n"
                "Install with:\n"
                "  sudo apt install python3-libzim\n"
                "  OR pip install libzim"
            )
            return
        
        output_path = self.output_var.get()
        title = self.title_var.get()
        
        try:
            self.status_var.set("Creating ZIM...")
            self._log(f"Creating: {output_path}")
            
            # Create ZIM
            creator = ZIMCreator(output_path, title)
            
            for file_path in self.files:
                doc = DocumentParser.parse_file(file_path)
                if doc:
                    creator.add_document(doc)
                    self._log(f"✓ {os.path.basename(file_path)}")
                else:
                    self._log(f"✗ Failed: {os.path.basename(file_path)}")
            
            zim_path = creator.create()
            
            size_mb = os.path.getsize(zim_path) / (1024 * 1024)
            self.status_var.set("Done!")
            self._log(f"\n✅ Created: {zim_path} ({size_mb:.1f} MB)")
            
            self.messagebox.showinfo(
                "Success!",
                f"ZIM file created successfully!\n\n"
                f"File: {zim_path}\n"
                f"Size: {size_mb:.1f} MB\n"
                f"Documents: {len(creator.documents)}\n\n"
                f"You can now use this with Hermit!"
            )
            
        except Exception as e:
            self.status_var.set("Error!")
            self._log(f"❌ Error: {e}")
            self.messagebox.showerror("Error", str(e))
    
    def run(self):
        """Start the GUI."""
        self.root.mainloop()


# =============================================================================
# CLI Interface
# =============================================================================

def cli_main():
    """Command-line interface."""
    parser = argparse.ArgumentParser(
        description="Forge - Create ZIM knowledge bases from documents",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  forge                          # Launch GUI
  forge /path/to/docs -o kb.zim  # Convert folder to ZIM
  forge file1.pdf file2.txt -o kb.zim  # Convert specific files
        """
    )
    
    parser.add_argument("inputs", nargs="*", help="Input files or directories")
    parser.add_argument("-o", "--output", default="knowledge.zim", help="Output ZIM file path")
    parser.add_argument("-t", "--title", default="My Knowledge Base", help="Knowledge base title")
    parser.add_argument("--gui", action="store_true", help="Force GUI mode")
    parser.add_argument("--no-recursive", action="store_true", help="Don't scan subdirectories")
    
    args = parser.parse_args()
    
    # GUI mode if no inputs
    if not args.inputs or args.gui:
        print("Launching Forge GUI...")
        gui = ForgeGUI()
        gui.run()
        return
    
    # CLI mode
    if not LIBZIM_AVAILABLE:
        print("ERROR: libzim not available.")
        print("Install with: sudo apt install python3-libzim")
        sys.exit(1)
    
    creator = ZIMCreator(args.output, args.title)
    
    for input_path in args.inputs:
        path = Path(input_path)
        
        if path.is_dir():
            count = creator.add_directory(str(path), recursive=not args.no_recursive)
            print(f"Added {count} documents from {path}")
        elif path.is_file():
            doc = DocumentParser.parse_file(str(path))
            if doc:
                creator.add_document(doc)
                print(f"Added: {path.name}")
        else:
            print(f"WARNING: Not found: {input_path}")
    
    if creator.documents:
        creator.create()
    else:
        print("ERROR: No documents were added.")
        sys.exit(1)


if __name__ == "__main__":
    cli_main()
